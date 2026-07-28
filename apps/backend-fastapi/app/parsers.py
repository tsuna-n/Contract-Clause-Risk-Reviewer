"""Document parsing: PDF/DOCX/TXT bytes -> normalized text + page offsets.

The offsets matter: every clause the pipeline produces carries a span back into
``ParsedDocument.text``, which is what makes citations checkable against the
source instead of trusting the model.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO

# --- data structures ---------------------------------------------------------


@dataclass
class TextSpan:
    """A run of text mapped back to its page and char offsets."""

    start: int
    end: int
    page: int


@dataclass
class ParsedDocument:
    """Normalized text plus the offset/page metadata needed for grounding.

    ``text`` is the normalized full-document string. ``spans`` map slices of
    ``text`` back to source pages. ``page_map`` maps a page number to the
    ``(start, end)`` char range it occupies in ``text``.
    """

    text: str
    spans: list[TextSpan] = field(default_factory=list)
    page_map: dict[int, tuple[int, int]] = field(default_factory=dict)

    def page_for_offset(self, offset: int) -> int | None:
        """Return the page number containing ``offset`` in ``text``."""
        for page, (start, end) in self.page_map.items():
            if start <= offset < end:
                return page
        return None


# --- normalization -----------------------------------------------------------

_LIGATURES = {"ﬁ": "fi", "ﬂ": "fl", "ﬀ": "ff"}

# Characters a Thai heading's title can begin with: the consonants ก–ฮ plus the
# four leading vowels เ แ โ ใ ไ, which are written *before* their consonant and
# so really do start words ("เงื่อนไข", "ใบแจ้งหนี้"). Thai has no case, so
# there is no uppercase test to make here - the script itself is the signal
# that a title, rather than more sentence, follows the number.
_THAI_LETTER = r"ก-ฮเ-ไ"
# Arabic digits plus Thai digits ๐–๙, spelled out rather than left to ``\d``:
# ``\d`` already matches every Unicode decimal digit, which would let Devanagari
# or Arabic-Indic numerals through in a rule that is deliberately about two
# scripts.
_DIGIT = r"0-9๐-๙"

#: Does a line open a numbered clause?
#:
#: ``1. Confidentiality`` / ``2.3 Term`` / ``12) Termination`` and their Thai
#: equivalents ``ข้อ 1. การรักษาความลับ``, ``๑. เงื่อนไข``. The optional
#: ``ข้อ``/``ข้อที่``/``ARTICLE``/``Section``/``Clause`` prefix is what most
#: real contracts put in front of the number.
#:
#: The trailing letter class is the part doing the real work: without it every
#: line starting with a figure ("500,000 baht payable...") reads as a heading.
_HEADING_RE = re.compile(
    r"^\s*"
    r"(?:(?:ข้อที่|ข้อ|ARTICLE|Article|Section|Clause)\s*)?"
    rf"([{_DIGIT}]+(?:\.[{_DIGIT}]+)*)"
    r"[.)]?\s+"
    rf"[A-Z{_THAI_LETTER}]"
)


def normalize_whitespace(text: str) -> str:
    """Collapse runs of whitespace while keeping paragraph breaks."""
    text = text.replace("\r\n", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def replace_ligatures(text: str) -> str:
    """Replace common typographic ligatures with ASCII equivalents."""
    for lig, repl in _LIGATURES.items():
        text = text.replace(lig, repl)
    return text


def is_heading(line: str) -> bool:
    """Heuristic: does ``line`` look like a numbered clause heading?

    Recognizes both scripts the corpus actually contains — ``1. Termination``
    and ``ข้อ 1. การเลิกสัญญา`` — because a document whose headings go
    unrecognized falls back to a blank-line paragraph split, and paragraph
    boundaries are not clause boundaries.
    """
    return bool(_HEADING_RE.match(line))


def normalize(text: str) -> str:
    """Apply the full normalization pipeline."""
    return normalize_whitespace(replace_ligatures(text))


# --- file formats ------------------------------------------------------------


def parse_pdf(data: bytes) -> ParsedDocument:
    """Parse PDF bytes into a :class:`ParsedDocument`, preserving page offsets."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype="pdf")
    try:
        raw_pages = [page.get_text() for page in doc]
    finally:
        doc.close()

    text_parts: list[str] = []
    spans: list[TextSpan] = []
    page_map: dict[int, tuple[int, int]] = {}
    offset = 0
    for page_number, raw_text in enumerate(raw_pages, start=1):
        page_text = normalize(raw_text)
        if not page_text:
            continue
        if text_parts:
            text_parts.append("\n\n")
            offset += 2
        start = offset
        text_parts.append(page_text)
        offset += len(page_text)
        spans.append(TextSpan(start=start, end=offset, page=page_number))
        page_map[page_number] = (start, offset)

    return ParsedDocument(text="".join(text_parts), spans=spans, page_map=page_map)


def parse_docx(data: bytes) -> ParsedDocument:
    """Parse DOCX bytes into a :class:`ParsedDocument`.

    DOCX has no real pages, so the whole document is a single synthetic page.
    """
    from docx import Document

    document = Document(BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    text = normalize("\n\n".join(paragraphs))
    span = TextSpan(start=0, end=len(text), page=1)
    return ParsedDocument(text=text, spans=[span], page_map={1: (0, len(text))})


#: Encodings tried, in order, when decoding a ``.txt`` upload.
#:
#: A plain-text file carries no encoding declaration, so this is a guess made in
#: the order the guesses are likely to be right. ``utf-8-sig`` covers plain UTF-8
#: as well, and strips the BOM Notepad writes; ``cp874`` is the Thai Windows
#: codepage, which is what a Thai contract saved as "ANSI" actually is - decoding
#: one as UTF-8 raises rather than mojibakes, so the fallback is reached.
_TXT_ENCODINGS = ("utf-8-sig", "cp874")


def _decode_text(data: bytes) -> str:
    """Decode text bytes, guessing the encoding, without ever raising.

    A file we cannot decode cleanly still reviews better mangled than not at
    all: the last resort replaces bad bytes instead of failing the upload.
    """
    for encoding in _TXT_ENCODINGS:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def parse_txt(data: bytes) -> ParsedDocument:
    """Parse plain-text bytes into a :class:`ParsedDocument`.

    Like DOCX, plain text has no pages, so the whole document is a single
    synthetic page. The segmenter still gets what it needs: ``normalize`` keeps
    blank-line paragraph breaks, and numbered headings survive as written.
    """
    text = normalize(_decode_text(data))
    span = TextSpan(start=0, end=len(text), page=1)
    return ParsedDocument(text=text, spans=[span], page_map={1: (0, len(text))})


#: Upload extension -> parser. ``ReviewService`` rejects anything not listed.
PARSERS = {"pdf": parse_pdf, "docx": parse_docx, "txt": parse_txt}

"""DOCX parsing via python-docx."""

from __future__ import annotations

from io import BytesIO

from app.parsers.models import ParsedDocument, TextSpan
from app.parsers.normalizer import normalize


def parse_docx(data: bytes) -> ParsedDocument:
    """Parse DOCX bytes into a :class:`ParsedDocument`.

    DOCX has no real pages, so the whole document is a single synthetic page.
    """
    from docx import Document

    document = Document(BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    text = normalize("\n\n".join(paragraphs))
    span = TextSpan(start=0, end=len(text), page=1)
    return ParsedDocument(text=text, spans=[span], page_map={1: (0, len(text))})

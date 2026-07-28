"""Unit tests for PDF/DOCX/TXT parsing."""

from io import BytesIO

from app.parsers import parse_docx, parse_pdf, parse_txt


def test_parse_pdf_extracts_text_and_page_map() -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "1. Confidentiality.")
    data = doc.write()
    doc.close()

    parsed = parse_pdf(data)
    assert "Confidentiality" in parsed.text
    assert parsed.page_map[1] == (0, len(parsed.text))
    assert parsed.page_for_offset(0) == 1


def test_parse_docx_joins_paragraphs_as_single_page() -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("1. Confidentiality.")
    document.add_paragraph("Each party shall keep information secret.")
    buf = BytesIO()
    document.save(buf)

    parsed = parse_docx(buf.getvalue())
    assert "Confidentiality" in parsed.text
    assert "secret" in parsed.text
    assert parsed.page_map == {1: (0, len(parsed.text))}


def test_parse_txt_keeps_paragraph_breaks_as_single_page() -> None:
    raw = "1. Confidentiality.\n\nEach party shall keep information secret.\n"

    parsed = parse_txt(raw.encode())
    assert "1. Confidentiality." in parsed.text
    # The blank line survives normalization: it is what the segmenter falls back
    # to when a document has no recognizable headings.
    assert "\n\n" in parsed.text
    assert parsed.page_map == {1: (0, len(parsed.text))}
    assert parsed.page_for_offset(0) == 1


def test_parse_txt_decodes_thai_in_utf8_and_cp874() -> None:
    thai = "ข้อ 1. การรักษาความลับ"

    assert parse_txt(thai.encode("utf-8")).text == thai
    # A Thai contract saved as "ANSI" from Windows is cp874, not UTF-8.
    assert parse_txt(thai.encode("cp874")).text == thai
    # And a UTF-8 file written by Notepad carries a BOM that must not survive
    # into the text, or the first heading stops matching.
    assert parse_txt(thai.encode("utf-8-sig")).text == thai


def test_parse_txt_never_raises_on_undecodable_bytes() -> None:
    parsed = parse_txt(b"\xff\xfe valid tail")
    assert "valid tail" in parsed.text

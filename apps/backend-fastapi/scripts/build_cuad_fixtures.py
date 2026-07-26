"""CLI: build the contract + gold fixtures from CUAD.

CUAD (the Contract Understanding Atticus Dataset) is 510 real commercial
contracts with expert-annotated spans across 41 review categories. This script
turns a checkout of it into the two fixtures the evaluation harness reads:

    data/contracts/<contract_id>.txt   the contract text, verbatim
    data/gold/annotations.jsonl        one gold record per contract

It also writes a few of them as .docx into data/samples/, because the review
endpoint only parses PDF and DOCX - without those there is nothing to drag into
the upload page and the app can't be exercised end to end.

Usage:
    python -m scripts.build_cuad_fixtures [--cuad ~/project/cuad] [--limit 12]

Why a generator instead of checked-in hand-written samples: the gold spans have
to line up with what :class:`~app.ai.agents.Segmenter` actually produces, and
with the character offsets that survive :func:`~app.parsers.normalize`. Both are
code, so deriving the fixtures from that same code is the only way they stay
correct when either changes. Re-run this script after touching the segmenter.

Three things about the labels, since they decide what the metrics mean:

* **Spans** are clause boundaries, not CUAD's highlights. CUAD marks the few
  lines that answer a category question ("...shall not exceed $50,000"); the
  reviewer works in whole clauses. So the document is segmented on its own
  numbered headings and each segment becomes one gold span - a complete
  segmentation, which is what ``segmentation_f1`` needs to be meaningful.
* **clause_type** comes from the CUAD annotation whose highlight falls inside
  the segment, mapped through :data:`CATEGORY_CLAUSE_TYPES`. Segments CUAD did
  not annotate are emitted *without* a type: CUAD's 41 categories don't cover
  the whole taxonomy (there is no confidentiality or force-majeure category),
  so labelling those segments ourselves would score the pipeline against
  guesses. ``run_eval`` counts an unlabelled segment for segmentation and skips
  it for classification/risk.
* **risk_level** is a policy call, not data: CUAD says a clause *is* a
  liability cap, not whether that is acceptable. :data:`CATEGORY_RISK` records
  the risk appetite the playbook encodes, from the perspective of the party
  running the review. Change it there and the playbook together.
"""

from __future__ import annotations

import argparse
import bisect
import json
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path

from app.ai.agents import Segmenter
from app.parsers import ParsedDocument, TextSpan, normalize
from app.schemas import ClauseType, RiskLevel

# --- CUAD category -> taxonomy ------------------------------------------------

#: CUAD category -> the taxonomy label in :class:`~app.schemas.ClauseType`.
#:
#: Four CUAD categories are deliberately absent - Document Name, Parties,
#: Agreement Date and Effective Date. Those are contract metadata that happens
#: to live on the cover page, not clauses a reviewer assesses risk on, and
#: including them would label a title block as a clause.
CATEGORY_CLAUSE_TYPES: dict[str, ClauseType] = {
    # Term and termination
    "Expiration Date": ClauseType.TERMINATION,
    "Renewal Term": ClauseType.TERMINATION,
    "Notice Period To Terminate Renewal": ClauseType.TERMINATION,
    "Termination For Convenience": ClauseType.TERMINATION,
    "Post-Termination Services": ClauseType.TERMINATION,
    # Governing law
    "Governing Law": ClauseType.GOVERNING_LAW,
    # Restrictive covenants
    "Non-Compete": ClauseType.NON_COMPETE,
    "Exclusivity": ClauseType.NON_COMPETE,
    "No-Solicit Of Customers": ClauseType.NON_COMPETE,
    "No-Solicit Of Employees": ClauseType.NON_COMPETE,
    "Non-Disparagement": ClauseType.NON_COMPETE,
    "Competitive Restriction Exception": ClauseType.NON_COMPETE,
    # Commercial / payment
    "Most Favored Nation": ClauseType.PAYMENT_TERMS,
    "Revenue/Profit Sharing": ClauseType.PAYMENT_TERMS,
    "Price Restrictions": ClauseType.PAYMENT_TERMS,
    "Minimum Commitment": ClauseType.PAYMENT_TERMS,
    "Volume Restriction": ClauseType.PAYMENT_TERMS,
    # Intellectual property
    "Ip Ownership Assignment": ClauseType.INTELLECTUAL_PROPERTY,
    "Joint Ip Ownership": ClauseType.INTELLECTUAL_PROPERTY,
    "License Grant": ClauseType.INTELLECTUAL_PROPERTY,
    "Non-Transferable License": ClauseType.INTELLECTUAL_PROPERTY,
    "Affiliate License-Licensor": ClauseType.INTELLECTUAL_PROPERTY,
    "Affiliate License-Licensee": ClauseType.INTELLECTUAL_PROPERTY,
    "Unlimited/All-You-Can-Eat-License": ClauseType.INTELLECTUAL_PROPERTY,
    "Irrevocable Or Perpetual License": ClauseType.INTELLECTUAL_PROPERTY,
    "Source Code Escrow": ClauseType.INTELLECTUAL_PROPERTY,
    # Liability
    "Uncapped Liability": ClauseType.LIMITATION_OF_LIABILITY,
    "Cap On Liability": ClauseType.LIMITATION_OF_LIABILITY,
    "Liquidated Damages": ClauseType.LIMITATION_OF_LIABILITY,
    # Warranty
    "Warranty Duration": ClauseType.WARRANTY,
    # No taxonomy label of their own - real clauses, but the taxonomy has no
    # bucket for assignment, audit or insurance, and OTHER is the honest answer
    # rather than filing them under a neighbouring type.
    "Rofr/Rofo/Rofn": ClauseType.OTHER,
    "Change Of Control": ClauseType.OTHER,
    "Anti-Assignment": ClauseType.OTHER,
    "Audit Rights": ClauseType.OTHER,
    "Insurance": ClauseType.OTHER,
    "Covenant Not To Sue": ClauseType.OTHER,
    "Third Party Beneficiary": ClauseType.OTHER,
}

#: CUAD category -> the risk the playbook assigns a clause of that kind.
#:
#: Read from the perspective of the party running the review. HIGH is a term
#: that transfers open-ended exposure or locks up the business (uncapped
#: liability, IP assigned away, an exclusivity or minimum-volume commitment);
#: LOW is a term that protects them or is pure housekeeping (a liability cap,
#: an escrow, a carve-out from a restrictive covenant).
CATEGORY_RISK: dict[str, RiskLevel] = {
    # Term and termination - being unable to exit is the risk, not the exit.
    "Expiration Date": RiskLevel.LOW,
    "Renewal Term": RiskLevel.MEDIUM,
    "Notice Period To Terminate Renewal": RiskLevel.MEDIUM,
    "Termination For Convenience": RiskLevel.MEDIUM,
    "Post-Termination Services": RiskLevel.MEDIUM,
    "Governing Law": RiskLevel.LOW,
    # Restrictive covenants - these bind the business itself.
    "Non-Compete": RiskLevel.HIGH,
    "Exclusivity": RiskLevel.HIGH,
    "No-Solicit Of Customers": RiskLevel.MEDIUM,
    "No-Solicit Of Employees": RiskLevel.MEDIUM,
    "Non-Disparagement": RiskLevel.LOW,
    "Competitive Restriction Exception": RiskLevel.LOW,  # a carve-out, i.e. relief
    # Commercial - open-ended pricing or volume promises.
    "Most Favored Nation": RiskLevel.HIGH,
    "Minimum Commitment": RiskLevel.HIGH,
    "Revenue/Profit Sharing": RiskLevel.MEDIUM,
    "Price Restrictions": RiskLevel.MEDIUM,
    "Volume Restriction": RiskLevel.MEDIUM,
    # Intellectual property - giving ownership or an unbounded licence away.
    "Ip Ownership Assignment": RiskLevel.HIGH,
    "Joint Ip Ownership": RiskLevel.HIGH,
    "Unlimited/All-You-Can-Eat-License": RiskLevel.HIGH,
    "Irrevocable Or Perpetual License": RiskLevel.HIGH,
    "License Grant": RiskLevel.MEDIUM,
    "Non-Transferable License": RiskLevel.MEDIUM,
    "Affiliate License-Licensor": RiskLevel.MEDIUM,
    "Affiliate License-Licensee": RiskLevel.MEDIUM,
    "Source Code Escrow": RiskLevel.LOW,
    # Liability
    "Uncapped Liability": RiskLevel.HIGH,
    "Liquidated Damages": RiskLevel.HIGH,
    "Cap On Liability": RiskLevel.LOW,
    # Everything else
    "Warranty Duration": RiskLevel.MEDIUM,
    "Rofr/Rofo/Rofn": RiskLevel.MEDIUM,
    "Change Of Control": RiskLevel.MEDIUM,
    "Anti-Assignment": RiskLevel.MEDIUM,
    "Audit Rights": RiskLevel.MEDIUM,
    "Insurance": RiskLevel.MEDIUM,
    "Covenant Not To Sue": RiskLevel.MEDIUM,
    "Third Party Beneficiary": RiskLevel.LOW,
}

#: Worst-first, for picking a winner when one clause carries several categories.
_RISK_ORDER = [RiskLevel.HIGH, RiskLevel.MEDIUM, RiskLevel.LOW, RiskLevel.UNKNOWN]


# --- offset-preserving normalization -----------------------------------------
#
# ``normalize`` rewrites the text (ligatures, collapsed runs of whitespace), so
# a CUAD offset into the raw contract does not point at the same character
# afterwards. The eval harness scores against the *normalized* text, so every
# annotation has to be carried across that rewrite rather than used as-is.


def _sub_with_map(
    pattern: str, repl: str, text: str, index_map: list[int]
) -> tuple[str, list[int]]:
    """Apply one regex substitution, carrying the source index of each char.

    ``index_map[i]`` is the offset in the *original* raw text that produced
    ``text[i]``. Replacement characters inherit the offset of the match start,
    which is what makes a collapsed run of spaces still point at where the run
    began.
    """
    out: list[str] = []
    out_map: list[int] = []
    pos = 0
    for match in re.finditer(pattern, text):
        out.extend(text[pos : match.start()])
        out_map.extend(index_map[pos : match.start()])
        out.extend(repl)
        out_map.extend([index_map[match.start()]] * len(repl))
        pos = match.end()
    out.extend(text[pos:])
    out_map.extend(index_map[pos:])
    return "".join(out), out_map


def normalize_with_offsets(text: str) -> tuple[str, list[int]]:
    """Return ``normalize(text)`` plus a raw-offset per normalized character.

    Mirrors :func:`app.parsers.normalize` step for step. It is duplicated
    rather than imported because the parser has no reason to carry index maps
    at runtime - only fixture generation needs to invert the transform.
    :func:`build_record` checks the two agree on every contract it writes, so
    this copy cannot silently drift from the parser.
    """
    index_map = list(range(len(text)))
    # replace_ligatures - one char becomes two, so offsets shift from here on.
    for ligature, replacement in (("ﬁ", "fi"), ("ﬂ", "fl"), ("ﬀ", "ff")):
        text, index_map = _sub_with_map(re.escape(ligature), replacement, text, index_map)
    # normalize_whitespace
    text, index_map = _sub_with_map(r"\r\n", "\n", text, index_map)
    text, index_map = _sub_with_map(r"[ \t]+", " ", text, index_map)
    text, index_map = _sub_with_map(r"\n{3,}", "\n\n", text, index_map)
    # .strip()
    start = len(text) - len(text.lstrip())
    end = len(text.rstrip())
    return text[start:end], index_map[start:end]


def raw_to_normalized(index_map: list[int], raw_offset: int) -> int | None:
    """Map a raw-text offset to the first normalized offset at or after it.

    ``index_map`` is non-decreasing, so this is a binary search. Returns
    ``None`` when the offset falls past the end - i.e. inside trailing
    whitespace that ``strip()`` removed.
    """
    position = bisect.bisect_left(index_map, raw_offset)
    return position if position < len(index_map) else None


# --- selection ---------------------------------------------------------------


@dataclass(frozen=True)
class Candidate:
    """One CUAD contract, with the stats the selection rule reads."""

    title: str
    text: str
    #: category -> raw-text offsets where the annotator highlighted an answer.
    annotations: dict[str, list[int]]
    heading_count: int

    @property
    def clause_types(self) -> set[ClauseType]:
        """Taxonomy labels this contract can contribute to the gold set."""
        return {
            CATEGORY_CLAUSE_TYPES[category]
            for category in self.annotations
            if category in CATEGORY_CLAUSE_TYPES
        }


#: Selection bounds. Too short and the "contract" is a one-page filing stub with
#: no clause structure; too long and a single eval run costs hundreds of LLM
#: calls. ``MIN_HEADINGS`` keeps documents the segmenter can actually split on
#: headings rather than falling back to paragraph splitting.
MIN_CHARS, MAX_CHARS = 8_000, 20_000
MIN_HEADINGS = 8
MIN_CATEGORIES = 8


def load_candidates(cuad_dir: Path) -> list[Candidate]:
    """Read CUADv1.json out of ``data.zip`` and return the eligible contracts."""
    archive = cuad_dir / "data.zip"
    if not archive.exists():
        raise SystemExit(f"CUAD archive not found: {archive}")

    with zipfile.ZipFile(archive) as zf, zf.open("CUADv1.json") as handle:
        dataset = json.load(handle)

    candidates: list[Candidate] = []
    for document in dataset["data"]:
        paragraph = document["paragraphs"][0]
        text = paragraph["context"]
        if not MIN_CHARS <= len(text) <= MAX_CHARS:
            continue

        normalized, _ = normalize_with_offsets(text)
        heading_count = sum(1 for _ in Segmenter._heading_boundaries(normalized))
        if heading_count < MIN_HEADINGS:
            continue

        annotations: dict[str, list[int]] = {}
        for qa in paragraph["qas"]:
            if not qa["answers"]:
                continue
            # CUAD ids are "<title>__<Category>".
            category = qa["id"].rsplit("__", 1)[-1]
            annotations[category] = [answer["answer_start"] for answer in qa["answers"]]
        if len(annotations) < MIN_CATEGORIES:
            continue

        candidates.append(
            Candidate(
                title=document["title"],
                text=text,
                annotations=annotations,
                heading_count=heading_count,
            )
        )
    return candidates


def select(candidates: list[Candidate], limit: int) -> list[Candidate]:
    """Pick ``limit`` contracts, greedily maximising taxonomy coverage.

    Each round takes the contract adding the most clause types not yet seen,
    breaking ties on annotation count and then title so the output is the same
    on every run. Coverage is the point: a gold set of ten distribution
    agreements would measure one corner of the taxonomy very precisely and the
    rest not at all.
    """
    remaining = sorted(candidates, key=lambda c: (-len(c.annotations), c.title))
    chosen: list[Candidate] = []
    covered: set[ClauseType] = set()

    while remaining and len(chosen) < limit:
        best = max(remaining, key=lambda c: len(c.clause_types - covered))
        chosen.append(best)
        covered |= best.clause_types
        remaining.remove(best)

    return chosen


# --- fixture writing ---------------------------------------------------------

_SLUG_STRIP = re.compile(r"[^a-z0-9]+")
#: The trailing plain-English phrase in a CUAD title ("...-DISTRIBUTOR
#: AGREEMENT"). Everything before it is filing machinery - dates, form types,
#: exhibit numbers - which all use characters this class excludes, so the match
#: starts exactly where the readable name does. A trailing digit is CUAD's
#: own disambiguator for a contract split across several excerpts.
_TITLE_TAIL = re.compile(r"([A-Za-z][A-Za-z ,&.]{3,})\d*$")


def contract_id(title: str) -> str:
    """Turn a CUAD title into a short, stable, filesystem-safe id.

    CUAD titles are filing artefacts - "GluMobileInc_20070319_S-1A_EX-10.09_
    436630_EX-10.09_Content License Agreement4". The company and the agreement
    type are what a reviewer recognises, so keep those two and drop everything
    in between.
    """
    company = _SLUG_STRIP.sub("-", title.split("_")[0].lower()).strip("-")
    tail = _TITLE_TAIL.search(title)
    kind = _SLUG_STRIP.sub("-", tail.group(1).lower()).strip("-") if tail else ""
    # Titles with no filing prefix at all ("VIVINT SOLAR, INC. - NON-COMPETITION
    # AGREEMENT") already carry the agreement type inside `company`; appending
    # it again would only repeat it.
    if not kind or kind in company:
        return company
    return f"{company}-{kind}"


def _pick_label(categories: list[str]) -> tuple[ClauseType, RiskLevel]:
    """Resolve the several CUAD categories inside one clause to one label.

    A "4. Term and Termination" clause typically carries three termination
    categories and one stray pricing one, so the type with the most
    annotations wins - counting categories, not picking the riskiest one,
    is what keeps that clause labelled ``termination``.

    Risk is then the worst risk within the winning type: a clause that caps
    liability and then carves an uncapped exception out of the cap is an
    uncapped-liability clause as far as the reviewer is concerned. Ties break
    towards the riskier type and then alphabetically, so the label never
    depends on dict ordering.
    """
    by_type: dict[ClauseType, list[str]] = {}
    for category in categories:
        by_type.setdefault(CATEGORY_CLAUSE_TYPES[category], []).append(category)

    def worst_risk(members: list[str]) -> RiskLevel:
        return min((CATEGORY_RISK[c] for c in members), key=_RISK_ORDER.index)

    clause_type = min(
        by_type,
        key=lambda t: (-len(by_type[t]), _RISK_ORDER.index(worst_risk(by_type[t])), t.value),
    )
    return clause_type, worst_risk(by_type[clause_type])


def build_record(candidate: Candidate) -> dict:
    """Build the gold record for one contract."""
    normalized, index_map = normalize_with_offsets(candidate.text)
    if normalized != normalize(candidate.text):
        raise SystemExit(
            f"normalize_with_offsets has drifted from app.parsers.normalize "
            f"on {candidate.title!r}; fix it before the offsets are trusted"
        )
    document = ParsedDocument(
        text=normalized,
        spans=[TextSpan(start=0, end=len(normalized), page=1)],
        page_map={1: (0, len(normalized))},
    )
    # The real segmenter, not a copy of its rules: gold spans that the pipeline
    # cannot reproduce would make segmentation_f1 measure the fixture, not the
    # pipeline. Segmenter.run never touches the LLM, so ``None`` is enough.
    clauses = Segmenter(None).run(document)  # type: ignore[arg-type]

    starts = [clause.span.start for clause in clauses]
    per_clause: dict[int, list[str]] = {}
    for category, raw_offsets in candidate.annotations.items():
        if category not in CATEGORY_CLAUSE_TYPES:
            continue
        for raw_offset in raw_offsets:
            offset = raw_to_normalized(index_map, raw_offset)
            if offset is None:
                continue
            # bisect_right - 1 is the clause whose span contains the offset.
            index = bisect.bisect_right(starts, offset) - 1
            if index >= 0:
                per_clause.setdefault(index, []).append(category)

    gold_clauses = []
    for index, clause in enumerate(clauses):
        entry: dict = {"span": {"start": clause.span.start, "end": clause.span.end}}
        categories = per_clause.get(index)
        if categories:
            clause_type, risk = _pick_label(categories)
            entry["clause_type"] = clause_type.value
            entry["risk_level"] = risk.value
            entry["cuad_categories"] = sorted(set(categories))
        gold_clauses.append(entry)

    return {
        "contract_id": contract_id(candidate.title),
        "source": "CUAD v1 (Atticus Project, CC BY 4.0)",
        "source_title": candidate.title,
        "clauses": gold_clauses,
    }


def write_docx(text: str, path: Path) -> None:
    """Write ``text`` as a .docx one paragraph per blank-line-separated block.

    ``parse_docx`` reads paragraphs and rejoins them with blank lines, so
    keeping the source's paragraph breaks is what makes the uploaded document
    segment into the same clauses the .txt fixture does.
    """
    from docx import Document

    document = Document()
    for block in re.split(r"\n\s*\n", text.strip()):
        document.add_paragraph(block.strip())
    document.save(path)


def main(argv: list[str]) -> int:
    """Regenerate ``data/contracts`` and ``data/gold/annotations.jsonl``."""
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--cuad",
        type=Path,
        default=Path.home() / "project" / "cuad",
        help="checkout of the CUAD repository (must contain data.zip)",
    )
    parser.add_argument("--limit", type=int, default=12, help="how many contracts to keep")
    parser.add_argument(
        "--samples",
        type=int,
        default=3,
        help="how many of them to also write as uploadable .docx",
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=Path("data"),
        help="the backend's data/ directory",
    )
    args = parser.parse_args(argv[1:])

    candidates = load_candidates(args.cuad)
    if not candidates:
        raise SystemExit("no CUAD contract met the selection criteria")
    selected = select(candidates, args.limit)

    contracts_dir = args.out / "contracts"
    contracts_dir.mkdir(parents=True, exist_ok=True)
    # Previous fixtures are removed, not merged: a stale .txt whose offsets no
    # longer match any gold record is worse than a missing one, because the
    # harness would happily score against it.
    for stale in contracts_dir.glob("*.txt"):
        stale.unlink()

    records = []
    for candidate in selected:
        record = build_record(candidate)
        (contracts_dir / f"{record['contract_id']}.txt").write_text(candidate.text)
        records.append(record)

    gold_path = args.out / "gold" / "annotations.jsonl"
    gold_path.parent.mkdir(parents=True, exist_ok=True)
    gold_path.write_text("".join(json.dumps(record) + "\n" for record in records))

    # The shortest ones, so a demo upload finishes in minutes rather than an
    # hour - the pipeline makes roughly four LLM calls per clause.
    samples_dir = args.out / "samples"
    samples_dir.mkdir(parents=True, exist_ok=True)
    for stale in samples_dir.glob("*.docx"):
        stale.unlink()
    shortest = sorted(zip(records, selected, strict=True), key=lambda pair: len(pair[1].text))
    for record, candidate in shortest[: args.samples]:
        write_docx(candidate.text, samples_dir / f"{record['contract_id']}.docx")

    labelled = sum(
        1 for record in records for clause in record["clauses"] if "clause_type" in clause
    )
    total = sum(len(record["clauses"]) for record in records)
    covered = sorted(
        {
            clause["clause_type"]
            for record in records
            for clause in record["clauses"]
            if "clause_type" in clause
        }
    )
    print(f"[build_cuad_fixtures] {len(records)} contracts -> {contracts_dir}")
    print(f"[build_cuad_fixtures] {args.samples} uploadable .docx -> {samples_dir}")
    print(f"[build_cuad_fixtures] {labelled}/{total} clauses carry a CUAD label")
    print(f"[build_cuad_fixtures] clause types covered: {', '.join(covered)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
